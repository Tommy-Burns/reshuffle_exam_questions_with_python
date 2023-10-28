from docx import Document

# Create a new Document
doc = Document()

# Add the questions and options to the document
questions = [
    ("Which planet in our solar system is known as the 'Red Planet'?", ["a) Jupiter", "b) Mars*", "c) Saturn", "d) Venus"]),
    ("What is the name of the galaxy that contains our solar system?", ["a) Centaurus A", "b) Milky Way*", "c) Triangulum", "d) Andromeda"]),
    ("What phenomenon occurs when a planet passes between Earth and the Sun, appearing to move across the Sun's disk?", ["a) Umbra", "b) Transit*", "c) Solar eclipse", "d) Solar flare"]),
    ("What is a light-year a measure of?", ["a) Brightness", "b) Distance*", "c) Temperature", "d) Time"]),
    ("Which of the following is not a type of telescope?", ["a) Radiospectrometer*", "b) Newtonian", "c) Cassegrain", "d) Refractor"]),
    ("What causes the changing phases of the Moon as observed from Earth?", ["a) The Moon's varying distance from Earth", "b) Shadows cast by Earth onto the Moon", "c) Sunlight falling on different parts of the Moon as it orbits Earth*", "d) The Moon's rotation on its axis"])
]

# Loop through the questions and add them to the document
for i, (question, options) in enumerate(questions, start=1):
    doc.add_paragraph(f"Question {i}) {question}")
    for option in options:
        doc.add_paragraph(option)

# Save the document to a file
doc.save("quiz_questions.docx")

print("Quiz questions have been successfully written to quiz_questions.docx.")
